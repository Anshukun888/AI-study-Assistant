"""
PDF and image processing: PyMuPDF (fitz), optional OCR for images.
"""
import fitz  # PyMuPDF
import os
import json
from typing import Optional, List, Dict, Tuple
from fastapi import UploadFile, HTTPException, status


def extract_text_from_image_bytes(data: bytes, filename: str = "") -> str:
    """Extract text from image using OCR (pytesseract) if available."""
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(data))
        if img.mode not in ("L", "RGB", "RGBA"):
            img = img.convert("RGB")
        try:
            import pytesseract
            return pytesseract.image_to_string(img) or "(No text detected in image)"
        except ImportError:
            return "[Image received. Install pytesseract for OCR text extraction.]"
    except Exception as e:
        return f"[Could not process image: {e}]"


def extract_text_from_docx_bytes(data: bytes) -> Optional[str]:
    """Extract text from DOCX file bytes. Returns None if python-docx is not installed."""
    try:
        from docx import Document as DocxDocument
        import io
        doc = DocxDocument(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        tables_text.append(cell.text.strip())
        text = "\n\n".join(paragraphs)
        if tables_text:
            text += "\n\n" + "\n".join(tables_text)
        return text.strip() or "(No text in document)"
    except ImportError:
        return None
    except Exception as e:
        return f"[Could not process DOCX: {e}]"


async def extract_text_from_pdf(file: UploadFile, max_size_mb: int = 50) -> str:
    """
    Extract text from uploaded PDF file.
    
    Args:
        file: Uploaded file object
        max_size_mb: Maximum file size in MB
        
    Returns:
        Extracted text as string
        
    Raises:
        HTTPException: If file is too large or invalid
    """
    # Check file size
    file_content = await file.read()
    file_size_mb = len(file_content) / (1024 * 1024)
    
    if file_size_mb > max_size_mb:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size ({file_size_mb:.2f} MB) exceeds maximum allowed size ({max_size_mb} MB)"
        )
    
    # Validate PDF
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File must be a PDF"
        )
    
    try:
        # Open PDF from memory
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        
        # Extract text from each page
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
        
        pdf_document.close()
        
        # Combine all text
        full_text = "\n\n".join(text_parts)
        
        if not full_text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="PDF contains no extractable text"
            )
        
        return full_text
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Error processing PDF: {str(e)}"
        )


def chunk_text(text: str, chunk_size: int = 3000, overlap: int = 200) -> list[str]:
    """
    Split text into chunks for processing by AI model.
    
    Args:
        text: Text to chunk
        chunk_size: Maximum size of each chunk in characters
        overlap: Number of characters to overlap between chunks
        
    Returns:
        List of text chunks
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        # Try to break at sentence boundary
        if end < len(text):
            # Look for sentence endings
            for break_char in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                last_break = text.rfind(break_char, start, end)
                if last_break != -1:
                    end = last_break + len(break_char)
                    break
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start with overlap
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks


async def extract_text_with_pages(file: UploadFile, max_size_mb: int = 50) -> Tuple[str, Dict[int, str]]:
    """
    Extract text from PDF with page number tracking.
    
    Returns:
        Tuple of (full_text, page_texts_dict) where page_texts_dict maps page_num -> text
    """
    file_content = await file.read()
    file_size_mb = len(file_content) / (1024 * 1024)
    
    if file_size_mb > max_size_mb:
        raise HTTPException(
            status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
            detail=f"File size ({file_size_mb:.2f} MB) exceeds maximum allowed size ({max_size_mb} MB)"
        )
    
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="File must be a PDF"
        )
    
    try:
        pdf_document = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        page_texts = {}
        
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            text = page.get_text()
            if text.strip():
                page_texts[page_num + 1] = text  # 1-indexed page numbers
                text_parts.append(text)
        
        pdf_document.close()
        
        full_text = "\n\n".join(text_parts)
        
        if not full_text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="PDF contains no extractable text"
            )
        
        return full_text, page_texts
        
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Error processing PDF: {str(e)}"
        )


def chunk_text_with_pages(text: str, page_texts: Dict[int, str], chunk_size: int = 3000, overlap: int = 200) -> List[Dict[str, any]]:
    """
    Split text into chunks with page number tracking.
    
    Returns:
        List of dicts with keys: 'text', 'pages' (list of page numbers)
    """
    if len(text) <= chunk_size:
        # Find which pages this text spans
        pages = []
        for page_num, page_text in page_texts.items():
            if page_text in text or text in page_text:
                pages.append(page_num)
        return [{'text': text, 'pages': sorted(set(pages)) if pages else [1]}]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        if end < len(text):
            for break_char in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                last_break = text.rfind(break_char, start, end)
                if last_break != -1:
                    end = last_break + len(break_char)
                    break
        
        chunk_text = text[start:end].strip()
        if chunk_text:
            # Find pages for this chunk
            pages = []
            for page_num, page_text in page_texts.items():
                # Check if chunk overlaps with page text
                chunk_words = set(chunk_text.lower().split()[:10])  # First 10 words
                page_words = set(page_text.lower().split()[:10])
                if chunk_words & page_words:  # Intersection
                    pages.append(page_num)
            
            chunks.append({
                'text': chunk_text,
                'pages': sorted(set(pages)) if pages else [1]
            })
        
        start = end - overlap
        if start >= len(text):
            break
    
    return chunks


def find_page_for_text(query_text: str, page_texts: Dict[int, str]) -> List[int]:
    """
    Find page numbers containing the query text.
    
    Returns:
        List of page numbers (1-indexed)
    """
    query_lower = query_text.lower()
    query_words = set(query_lower.split())
    matching_pages = []
    
    for page_num, page_text in page_texts.items():
        page_lower = page_text.lower()
        # Check if significant words match
        if len(query_words) > 0:
            matches = sum(1 for word in query_words if word in page_lower)
            if matches >= min(2, len(query_words) * 0.3):  # At least 30% match or 2 words
                matching_pages.append(page_num)
    
    return sorted(set(matching_pages)) if matching_pages else []

